# This script creates reports for clients

import sys
from docx import Document
from docx.shared import Inches
import pandas as pd
import pandas
import datetime
import unidecode

# Script creates word dco reports.

# Strings for report fo find mutations

leiden=("Koagulační faktor V", "F5; G1691A (Leiden)", "F5; NM_000130.4; c.1601G>A")
protrombin=("Protrombin", "F2; G20210A", "F2; NM_000506.4; c.*97G>A")
MTHFR1 = ("Methyltetrahydrofolát reduktáza", "MTHFR; C677T", "MTHFR; NM_005957.4; c.665C>T")
MTHFR2 = ("Methyltetrahydrofolát reduktáza", "MTHFR; A1298C", "MTHFR; NM_005957.4; c.1286A>C")


# Class for client information

class Client:
	def __init__(self, protocol_num, exam_type, name, date_of_exam, birth_num, isurance_comp_num):
		self.protocol_num = protocol_num
		self.exam_type = exam_type
		self.name = name
		self.date_of_exam = date_of_exam
		self.birth_num = birth_num
		self.isurance_comp_num = isurance_comp_num
		self.results_text = ()	
		
	def set_genotype_text(self, text):
		self.results_text = self.results_text+text
	
	def get_genotype_text(self):
		return self.results_text
		
# Class that find info about mutations and decide what is reportable
		
class Make_reports:
	def __init__(self, clients_list_file):
		self.clients = []
		self.clients_list_file = pandas.read_excel(clients_list_file)
		text = []
		t = []
	
	# read individual client peronal info and genotype info
	
	def read_clients_info(self):
		client_list = self.clients_list_file.values.tolist()
		for row in range(0, len(client_list)):
			self.clients.append(Client(client_list[row][0],client_list[row][1],client_list[row][2],client_list[row][3],client_list[row][4],client_list[row][5]))
		for client in self.clients:
			file_name = ("Data/"+client.protocol_num.replace('/', '_')+"_"+unidecode.unidecode(client.name.split()[1])+".out.csv")
			try:
				file = pd.read_csv(file_name, sep=',')
			except IOError:
				print('File error', file_name,  sys.stderr)
				sys.exit(1)
			self.get_genotype(file, client)
			Make_word_report(client)
			
			
	def get_reportable(self, client_data):
		return client_data.loc[((client_data['analysis'] == 'TROMBO_basic') | (client_data['analysis'] == 'TROMBO_full')) & (client_data['report'] == True)]
	
	# returns line that passes given conditions of mutations
	
	def get_mutated_chr_line(self, client_data, start, end, ref, alt):
		reportable_data = self.get_reportable(client_data)
		return reportable_data.loc[(reportable_data['Start'] == start) & (reportable_data['End'] == end) & (reportable_data['Ref'] == ref) & (reportable_data['Alt'] == alt)]
	
	def get_var_freq(self, mutated_chr_line):
		return pd.to_numeric(mutated_chr_line['variant_frequency'])
	
	# tests for types of genotype and constucts result text
	
	def get_genotype(self, file, client):
			client_data = pd.DataFrame(file)
			
			#test for F5; G1691A (Leiden)
			mutated_chr_line = self.get_mutated_chr_line(client_data, 169519049, 169519049, 'T', 'C')
			if not mutated_chr_line.empty:
				var_freq = float(self.get_var_freq(mutated_chr_line))
				if var_freq > 0.8:
					g = ('G/G(WT)',)
					client.set_genotype_text(leiden+g)
				elif var_freq < 0.2:
					g = ('A/A(HOM)',)
					client.set_genotype_text(leiden+g)
				else:
					g = ('G/A(HET)',)
					client.set_genotype_text(leiden+g)
					
			#test for F2; G20210A (protrombin)
			mutated_chr_line = self.get_mutated_chr_line(client_data, 46761055, 46761055, 'G', 'A')
			if not mutated_chr_line.empty:
				var_freq = float(self.get_var_freq(mutated_chr_line))
				if var_freq > 0.8:
					g = ('A/A(HOM)',)
					client.set_genotype_text(protrombin+g)
				elif var_freq < 0.2:
					g = ('G/G(WT)',)
					client.set_genotype_text(protrombin+g)
				else:
					g = ('G/A(HET)',)
					client.set_genotype_text(protrombin+g)
			
			#test for MTHFR; C677T
			mutated_chr_line = self.get_mutated_chr_line(client_data, 11856378, 11856378, 'G', 'A')
			if not mutated_chr_line.empty:
				var_freq = float(self.get_var_freq(mutated_chr_line))
				if var_freq > 0.8:
					g = ('T/T(HOM)',)
					client.set_genotype_text(MTHFR1+g)
				elif var_freq < 0.2:
					g = ('C/C(WT)',)
					client.set_genotype_text(MTHFR1+g)
				else:
					g = ('C/T(HET)',)
					client.set_genotype_text(MTHFR1+g)
			
			#test for MTHFR; A1298C
			mutated_chr_line = self.get_mutated_chr_line(client_data, 11854476, 11854476, 'T', 'G')
			if not mutated_chr_line.empty:
				var_freq = float(self.get_var_freq(mutated_chr_line))
				if var_freq > 0.8:
					g = ('C/C(HOM)',)
					client.set_genotype_text(MTHFR2+g)
				elif var_freq < 0.2:
					g = ('A/A(WT)',)
					client.set_genotype_text(MTHFR2+g)
				else:
					g = ('A/C(HET)',)
					client.set_genotype_text(MTHFR2+g)
					
#  makes report as word ducument
			
def Make_word_report(client):
	document = Document()

	document.add_picture('Ukazka/hlavicka.png', width=Inches(5.8))
	document.add_picture('Ukazka/repromeda_line.png', width=Inches(5.8))
	document.add_heading('\tVýsledný protokol molekulárně-genetického vyšetření trombofilních mutací metodou MPS', 2)

	records = (
	('Protokol č.:', client.protocol_num+'_T'),
	('Odesílající zdravotnické zařízení:', 'REPROMEDA s.r.o., MUDr. Ilga Grochová \nStudentská 812/6, 62500 Brno'),
	('Jméno a příjmení:', client.name),
	('Číslo pojištěnce:', client.birth_num),
	('Datum odběru:', client.date_of_exam.strftime("%d.%m.%Y")),
	('Popis vzorku:', 'Periferní krev'))
	
	# makes personal info table
	
	table1 = document.add_table(rows=0, cols=2)
	for i, j in records:
		row_cells = table1.add_row().cells
		row_cells[0].text = i
		row_cells[1].text = j

	document.add_paragraph('\nTabulka s výsledky:')
	table2 = document.add_table(rows=1, cols=4)
	
	hdr_cells = table2.rows[0].cells
	hdr_cells[0].text = 'Název rizikového faktoru'
	hdr_cells[1].text = 'Tradiční název varianty'
	hdr_cells[2].text = 'Gen; referenční sekvence; \nHGVS zápis varianty'
	hdr_cells[3].text = 'Genotyp'
	
	b=('Koagulační faktor V', 'F5; G1691A (Leiden)', 'F5; NM_000130.4; c.1601G>A', 'G/G(WT)', 'Methyltetrahydrofolát reduktáza', 'MTHFR; C677T', 'MTHFR; NM_005957.4; c.665C>T', 'C/T(HET)', 'Methyltetrahydrofolát reduktáza', 'MTHFR; A1298C', 'MTHFR; NM_005957.4; c.1286A>C', 'A/C(HET)')
	
	# makes result table
	
	for i in range(0, len(client.results_text),4):
		row_cells = table2.add_row().cells
		k = 0
		for j in range(i,i+4):
			row_cells[k].text = client.results_text[j]
			k = k + 1
	
	document.add_picture('Ukazka/repromeda_line.png', width=Inches(5.8))
	
	document.add_page_break()
	doc_name = ("Reporty/"+client.protocol_num.replace('/', '_')+"_"+unidecode.unidecode(client.name.split()[1])+"_TROMBO_CZE.docx")
	document.save(doc_name)
	
	# from docx2pdf package, does not work in linux, other methods of converting to pdf file should be used
	#convert(doc_name) #, "Reporty/"+client.protocol_num.replace('/', '_')+"_"+unidecode.unidecode(client.name.split()[1])+"_TROMBO_CZE.pdf")

def main():
	reports = Make_reports('Data/KLIENTI.xlsx')
	reports.read_clients_info()

if __name__ == '__main__':
    main()
